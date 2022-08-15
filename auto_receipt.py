#################################################
  ## Using a docx template, programmatically create PDFs company specific receipts of payment.
#################################################
from docx import Document
from docx.shared import Pt
from docx2pdf import convert
import numpy as np
import os
import pandas as pd

####### Bring in data:
data = pd.read_csv('template/payment_data.csv')
contact_info = pd.read_csv('template/contact_info.csv')

####### Add on & format company contact info:
data = pd.merge(data, contact_info, how='left', left_on='cluster_comm_uuid', right_on='agency_uuid')
data['state'] = data['state'].str.upper()
data['zip'] = data['zip'].astype(int).astype(str)
data['company_city_state_zip'] = data['city'] + ", " + data['state'] + " " + data['zip']

####### Subset to only elligible companies:
data = data[data['Eligibility'] == 1].reset_index(drop=True)


####### Format file titles:
docx_titles = data['company'] + '_' + data['state'].str.replace(' ', '_')


####### This is the list of the text to replace in the document
text_to_replace = [
    'REPLACE_COMPANY_NAME', 'ADDRESS_LINE_1', 'ADDRESS_LINE_2', 'COMPANY_CITY_STATE_ZIP',
    '***PAY_1', '***PAY_2']
####### Reduce data to only columns to replace data, in the same order as text_to_replace
data = data[['company_name', 'street_line_1', 'street_line_2', 'company_city_state_zip',
      'payment_1', 'payment_2']]
####### Change column names to the text that column will replace:
data.columns = text_to_replace




####### For each row in data, replace table & paragraph text when needed, then create doc
for agency_i in range(len(data)):
    ### Bring in new document each for row:
    document = Document('template/Final_Template.docx')
    for to_replace in text_to_replace:
        ### Replacing text inside a table, go through tables then cells one by one
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        ## If string to replace in cell then replace:
                        if to_replace in paragraph.text:
                            paragraph.text = paragraph.text.replace(to_replace, str(data[to_replace][agency_i]))
                            ### Font auto changes to 11 pt. Change font back to Arial 9 pt font:
                            font = paragraph.style.font
                            font.name = 'Arial'
                            font.size = Pt(9)

        ### Replacing non-table text, go thorugh paragraphs one by one
        for paragraph in document.paragraphs:
            ## If string to replace in cell then replace:
            if to_replace in paragraph.text:
                ## IF string is ADDRESS_LINE_2 and value is NaN (float type, all non-NaN are string) then delete line:
                if to_replace == 'ADDRESS_LINE_2' and type(data[to_replace][agency_i]) == float:
                    CT_P = paragraph._element
                    CT_P.getparent().remove(CT_P)
                else:
                    paragraph.text = str(data[to_replace][agency_i])

    ####### Save to docx file with company name 
    document.save('output/' + docx_titles[agency_i] + '.docx')




####### Convert docx files to pdf files:
convert("output/")

####### Remove .docx files
directory = 'output/'
files_in_directory = os.listdir(directory)
filtered_files = [file for file in files_in_directory if file.endswith(".docx")]
for file in filtered_files:
    path_to_file = os.path.join(directory, file)
    os.remove(path_to_file)
